"""
Unit tests for DocumentLoaderService — table formatting and loader routing.
"""

import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch

from app.services.document_loader_service import (
    DocumentLoaderService,
    docx_table_to_markdown,
    format_paragraph,
    table_to_markdown,
)

MEANINGFUL_PDF_TEXT = " ".join(
    [
        "This document contains enough meaningful alphabetic content to pass",
        "the native PDF extraction threshold without falling through to OCR.",
    ]
    * 8
)


class TestTableToMarkdown:
    def test_basic_table(self):
        table = [
            ["Name", "Score"],
            ["Alice", "95"],
            ["Bob", "87"],
        ]
        md = table_to_markdown(table)
        assert "| Name | Score |" in md
        assert "| Alice | 95 |" in md
        assert "| --- | --- |" in md

    def test_empty_table_returns_empty_string(self):
        assert table_to_markdown([]) == ""
        assert table_to_markdown(None) == ""

    def test_pads_short_rows(self):
        table = [["A", "B", "C"], ["1", "2"]]
        md = table_to_markdown(table)
        assert "| 1 | 2 |  |" in md


class TestFormatParagraph:
    def test_heading_styles(self):
        para = MagicMock()
        para.text = "Introduction"
        para.style.name = "Heading 1"
        assert format_paragraph(para) == "# Introduction"

        para.style.name = "Heading 3"
        assert format_paragraph(para) == "### Introduction"

    def test_list_style(self):
        para = MagicMock()
        para.text = "Item one"
        para.style.name = "List Paragraph"
        assert format_paragraph(para) == "- Item one"


class TestDocumentLoaderService:
    def test_load_text_file(self, tmp_path: Path):
        f = tmp_path / "note.txt"
        f.write_text("Hello world", encoding="utf-8")
        svc = DocumentLoaderService()
        assert svc.load(f, "text") == "Hello world"

    def test_load_pdf_uses_pymupdf4llm_first(self, tmp_path: Path):
        f = tmp_path / "doc.pdf"
        f.write_bytes(b"%PDF-1.4 fake")

        svc = DocumentLoaderService()
        with patch.object(
            svc, "_load_pdf_pymupdf4llm", return_value=f"# Title\n\n{MEANINGFUL_PDF_TEXT}"
        ) as mock_primary:
            result = svc.load(f, "pdf")
        mock_primary.assert_called_once()
        assert "Title" in result

    def test_load_pdf_falls_back_to_pdfplumber(self, tmp_path: Path):
        f = tmp_path / "doc.pdf"
        f.write_bytes(b"%PDF-1.4 fake")

        svc = DocumentLoaderService()
        with patch.object(svc, "_load_pdf_pymupdf4llm", side_effect=RuntimeError("fail")):
            with patch.object(
                svc,
                "_load_pdf_pdfplumber",
                return_value=f"## Page 1\n\nRecovered text. {MEANINGFUL_PDF_TEXT}",
            ) as mock_fallback:
                result = svc.load(f, "pdf")
        mock_fallback.assert_called_once()
        assert "Recovered text" in result

    def test_pdf_ocr_reports_missing_system_dependencies(self, tmp_path: Path):
        f = tmp_path / "scan.pdf"
        f.write_bytes(b"%PDF-1.4 fake")

        svc = DocumentLoaderService()
        with patch("app.services.document_loader_service.shutil.which", return_value=None):
            with pytest.raises(RuntimeError) as exc_info:
                svc._load_pdf_pytesseract(f)

        assert "Poppler and Tesseract" in str(exc_info.value)
        assert "pdfinfo" in str(exc_info.value)

    def test_load_docx_block_order(self, tmp_path: Path):
        pytest.importorskip("docx")
        from docx import Document

        f = tmp_path / "report.docx"
        doc = Document()
        doc.add_heading("Summary", level=1)
        doc.add_paragraph("Platform overview paragraph.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Metric"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Users"
        table.cell(1, 1).text = "100"
        doc.save(f)

        svc = DocumentLoaderService()
        text = svc.load(f, "docx")
        assert "Summary" in text
        assert "Platform overview" in text
        assert "| Metric | Value |" in text
        assert "| Users | 100 |" in text

    def test_docx_table_helper(self):
        mock_table = MagicMock()
        mock_row = MagicMock()
        mock_cell_a = MagicMock()
        mock_cell_a.text = "Col1"
        mock_cell_b = MagicMock()
        mock_cell_b.text = "Col2"
        mock_row.cells = [mock_cell_a, mock_cell_b]
        mock_table.rows = [mock_row]
        md = docx_table_to_markdown(mock_table)
        assert "| Col1 | Col2 |" in md
